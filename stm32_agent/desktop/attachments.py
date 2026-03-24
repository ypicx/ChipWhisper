from __future__ import annotations

import base64
import csv
import json
import mimetypes
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Sequence

try:
    import fitz
except Exception:  # pragma: no cover - optional dependency
    fitz = None

try:
    from docx import Document
except Exception:  # pragma: no cover - optional dependency
    Document = None

try:
    from openpyxl import load_workbook
except Exception:  # pragma: no cover - optional dependency
    load_workbook = None


ATTACHMENT_FILE_FILTER = (
    "Supported Files (*.pdf *.docx *.xlsx *.csv *.txt *.md *.json *.png *.jpg *.jpeg *.bmp *.webp);;"
    "Documents (*.pdf *.docx *.xlsx *.csv *.txt *.md *.json);;"
    "Images (*.png *.jpg *.jpeg *.bmp *.webp)"
)

_TEXT_SUFFIXES = {".txt", ".md", ".json"}
_TABLE_SUFFIXES = {".csv", ".xlsx"}
_DOCUMENT_SUFFIXES = {".pdf", ".docx"}
_IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".bmp", ".webp"}
_UNSUPPORTED_OFFICE_SUFFIXES = {".doc", ".xls"}
_MAX_TEXT_CHARS_PER_FILE = 6000
_MAX_IMAGE_BYTES = 4 * 1024 * 1024
_MAX_PDF_PAGES = 12
_MAX_TABLE_ROWS = 80
_MAX_TABLE_COLS = 12
_MAX_SHEETS = 5


@dataclass
class AttachmentDigest:
    path: str
    name: str
    suffix: str
    media_kind: str
    mime_type: str
    extracted_text: str = ""
    image_data_url: str = ""
    warnings: List[str] = field(default_factory=list)
    errors: List[str] = field(default_factory=list)

    def to_dict(self) -> dict[str, object]:
        return {
            "path": self.path,
            "name": self.name,
            "suffix": self.suffix,
            "media_kind": self.media_kind,
            "mime_type": self.mime_type,
            "extracted_text_preview": self.extracted_text[:300],
            "has_inline_image": bool(self.image_data_url),
            "warnings": self.warnings,
            "errors": self.errors,
        }

    def to_state_dict(self) -> dict[str, object]:
        return {
            "path": self.path,
            "name": self.name,
            "suffix": self.suffix,
            "media_kind": self.media_kind,
            "mime_type": self.mime_type,
            "extracted_text": self.extracted_text,
            "image_data_url": self.image_data_url,
            "warnings": list(self.warnings),
            "errors": list(self.errors),
        }

    @classmethod
    def from_state_dict(cls, payload: dict[str, object]) -> "AttachmentDigest":
        return cls(
            path=str(payload.get("path", "")).strip(),
            name=str(payload.get("name", "")).strip(),
            suffix=str(payload.get("suffix", "")).strip(),
            media_kind=str(payload.get("media_kind", "file")).strip() or "file",
            mime_type=str(payload.get("mime_type", "application/octet-stream")).strip() or "application/octet-stream",
            extracted_text=str(payload.get("extracted_text", "") or ""),
            image_data_url=str(payload.get("image_data_url", "") or ""),
            warnings=[str(item) for item in payload.get("warnings", []) if str(item).strip()],
            errors=[str(item) for item in payload.get("errors", []) if str(item).strip()],
        )

    def prompt_block(self) -> str:
        lines = [f"文件: {self.name}", f"类型: {self.media_kind}"]
        if self.extracted_text:
            lines.append("提取内容:")
            lines.append(self.extracted_text)
        elif self.image_data_url:
            lines.append("说明: 这是图片文件，若当前模型支持视觉，请结合图片本身理解需求。")
        else:
            lines.append("说明: 当前未提取出可直接送入模型的文本内容。")
        if self.warnings:
            lines.append("注意:")
            lines.extend(f"- {item}" for item in self.warnings)
        if self.errors:
            lines.append("错误:")
            lines.extend(f"- {item}" for item in self.errors)
        return "\n".join(lines)


@dataclass
class AttachmentBatch:
    attachments: List[AttachmentDigest]
    warnings: List[str]
    errors: List[str]

    def to_dict(self) -> dict[str, object]:
        return {
            "attachments": [item.to_dict() for item in self.attachments],
            "warnings": self.warnings,
            "errors": self.errors,
        }


def collect_attachment_digests(paths: Sequence[str | Path]) -> AttachmentBatch:
    attachments: List[AttachmentDigest] = []
    warnings: List[str] = []
    errors: List[str] = []
    seen: set[Path] = set()

    for raw_path in paths:
        path = Path(raw_path).expanduser()
        try:
            resolved = path.resolve()
        except OSError:
            resolved = path
        if resolved in seen:
            continue
        seen.add(resolved)

        if not resolved.exists():
            errors.append(f"附件不存在: {resolved}")
            continue
        if resolved.is_dir():
            errors.append(f"暂不支持直接添加目录: {resolved}")
            continue

        digest = _digest_file(resolved)
        attachments.append(digest)
        warnings.extend(digest.warnings)
        errors.extend(digest.errors)

    return AttachmentBatch(attachments=attachments, warnings=warnings, errors=errors)


def compose_multimodal_user_content(
    base_text: str,
    attachments: Sequence[AttachmentDigest] | None = None,
) -> str | list[dict[str, object]]:
    prompt = base_text.strip()
    normalized_attachments = [item for item in attachments or [] if item is not None]
    if not normalized_attachments:
        return prompt

    parts: list[dict[str, object]] = [{"type": "text", "text": prompt}]
    manifest_lines = [
        "以下是用户额外提供的需求附件，请结合它们一起理解需求。",
        "如果文档提取文本与图片信息冲突，以图片、原文件内容和用户原始意图为准。",
    ]
    for index, item in enumerate(normalized_attachments, start=1):
        manifest_lines.append("")
        manifest_lines.append(f"[附件 {index}]")
        manifest_lines.append(item.prompt_block())
    parts.append({"type": "text", "text": "\n".join(manifest_lines)})

    for item in normalized_attachments:
        if item.image_data_url:
            parts.append(
                {
                    "type": "image",
                    "image_url": item.image_data_url,
                    "detail": "auto",
                    "name": item.name,
                }
            )
    return parts


def render_attachment_list(paths: Sequence[str | Path]) -> str:
    items = [Path(path).name for path in paths if str(path).strip()]
    if not items:
        return ""
    return "\n".join(f"- {name}" for name in items)


def _digest_file(path: Path) -> AttachmentDigest:
    suffix = path.suffix.lower()
    mime_type = mimetypes.guess_type(path.name)[0] or "application/octet-stream"
    digest = AttachmentDigest(
        path=str(path),
        name=path.name,
        suffix=suffix,
        media_kind=_classify_kind(suffix),
        mime_type=mime_type,
    )

    try:
        if suffix in _TEXT_SUFFIXES:
            digest.extracted_text = _clip_text(_read_text_file(path))
        elif suffix == ".csv":
            digest.extracted_text = _clip_text(_read_csv_file(path))
        elif suffix == ".xlsx":
            digest.extracted_text = _clip_text(_read_xlsx_file(path, digest))
        elif suffix == ".pdf":
            digest.extracted_text = _clip_text(_read_pdf_file(path, digest))
        elif suffix == ".docx":
            digest.extracted_text = _clip_text(_read_docx_file(path, digest))
        elif suffix in _IMAGE_SUFFIXES:
            digest.image_data_url = _read_image_data_url(path, digest)
        elif suffix in _UNSUPPORTED_OFFICE_SUFFIXES:
            digest.warnings.append(f"暂不支持旧版 Office 格式: {path.suffix}，建议另存为 .docx 或 .xlsx。")
        else:
            digest.warnings.append(f"暂不支持该文件类型: {path.suffix or path.name}")
    except Exception as exc:  # pragma: no cover - defensive runtime guard
        digest.errors.append(f"{path.name} 解析失败: {exc}")

    if not digest.extracted_text and not digest.image_data_url and not digest.errors:
        digest.warnings.append(f"{path.name} 没有提取到可用内容。")
    return digest


def _classify_kind(suffix: str) -> str:
    if suffix in _TEXT_SUFFIXES:
        return "text"
    if suffix in _TABLE_SUFFIXES:
        return "table"
    if suffix in _DOCUMENT_SUFFIXES:
        return "document"
    if suffix in _IMAGE_SUFFIXES:
        return "image"
    return "file"


def _read_text_file(path: Path) -> str:
    encodings = ("utf-8", "utf-8-sig", "gb18030")
    last_error: Exception | None = None
    for encoding in encodings:
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError as exc:
            last_error = exc
            continue
    if last_error is not None:
        raise last_error
    return path.read_text(encoding="utf-8", errors="replace")


def _read_csv_file(path: Path) -> str:
    rows: list[str] = []
    with path.open("r", encoding="utf-8-sig", newline="", errors="replace") as handle:
        reader = csv.reader(handle)
        for row_index, row in enumerate(reader):
            if row_index >= _MAX_TABLE_ROWS:
                rows.append("... (CSV 已截断)")
                break
            cells = [str(cell).strip() for cell in row[:_MAX_TABLE_COLS]]
            rows.append(" | ".join(cells))
    return "\n".join(rows)


def _read_xlsx_file(path: Path, digest: AttachmentDigest) -> str:
    if load_workbook is None:
        digest.warnings.append("未安装 openpyxl，当前无法提取 .xlsx 内容。")
        return ""

    workbook = load_workbook(path, read_only=True, data_only=True)
    lines: list[str] = []
    try:
        for sheet_index, sheet_name in enumerate(workbook.sheetnames):
            if sheet_index >= _MAX_SHEETS:
                lines.append("... (工作表已截断)")
                break
            sheet = workbook[sheet_name]
            lines.append(f"[Sheet] {sheet_name}")
            for row_index, row in enumerate(sheet.iter_rows(values_only=True)):
                if row_index >= _MAX_TABLE_ROWS:
                    lines.append("... (当前工作表已截断)")
                    break
                values = ["" if cell is None else str(cell).strip() for cell in row[:_MAX_TABLE_COLS]]
                if any(values):
                    lines.append(" | ".join(values))
            lines.append("")
    finally:
        workbook.close()
    return "\n".join(line for line in lines if line is not None).strip()


def _read_pdf_file(path: Path, digest: AttachmentDigest) -> str:
    if fitz is None:
        digest.warnings.append("未安装 PyMuPDF，当前无法提取 PDF 文本。")
        return ""

    lines: list[str] = []
    with fitz.open(path) as document:
        total_pages = len(document)
        for page_index in range(min(total_pages, _MAX_PDF_PAGES)):
            page = document.load_page(page_index)
            text = page.get_text("text").strip()
            if not text:
                continue
            lines.append(f"[Page {page_index + 1}]")
            lines.append(text)
            lines.append("")
        if total_pages > _MAX_PDF_PAGES:
            digest.warnings.append(f"PDF 只提取了前 {_MAX_PDF_PAGES} 页。")
    return "\n".join(lines).strip()


def _read_docx_file(path: Path, digest: AttachmentDigest) -> str:
    if Document is None:
        digest.warnings.append("未安装 python-docx，当前无法提取 .docx 内容。")
        return ""

    document = Document(path)
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)
    for table_index, table in enumerate(document.tables, start=1):
        lines.append(f"[Table {table_index}]")
        for row in table.rows[:_MAX_TABLE_ROWS]:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells[:_MAX_TABLE_COLS]]
            if any(cells):
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def _read_image_data_url(path: Path, digest: AttachmentDigest) -> str:
    size = path.stat().st_size
    if size > _MAX_IMAGE_BYTES:
        digest.warnings.append(
            f"{path.name} 超过 {_MAX_IMAGE_BYTES // (1024 * 1024)}MB，当前不会内联发送给模型。"
        )
        return ""
    raw = path.read_bytes()
    encoded = base64.b64encode(raw).decode("ascii")
    return f"data:{digest.mime_type};base64,{encoded}"


def _clip_text(text: str) -> str:
    normalized = "\n".join(line.rstrip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n"))
    normalized = normalized.strip()
    if len(normalized) <= _MAX_TEXT_CHARS_PER_FILE:
        return normalized
    return normalized[:_MAX_TEXT_CHARS_PER_FILE].rstrip() + "\n... (内容已截断)"
